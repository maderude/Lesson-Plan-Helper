"""
ui/export.py — Export Helpers for Lesson Plan Helper
=====================================================

Converts a Markdown lesson plan into downloadable PDF and Word formats.
Also provides a copy-to-clipboard HTML snippet for the Streamlit UI.
"""

from __future__ import annotations

import io
import re
from typing import Any

# ==============================================================================
# Markdown → PDF  (via fpdf2)
# ==============================================================================

def markdown_to_pdf(md_text: str, meta: dict[str, Any] | None = None) -> bytes:
    """Convert a Markdown lesson plan to a styled PDF document.

    Parameters
    ----------
    md_text : str
        The lesson plan in Markdown format.
    meta : dict, optional
        Form metadata (subject, grade, etc.) for the PDF title.

    Returns
    -------
    bytes
        The PDF file content.
    """
    from fpdf import FPDF

    class LessonPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(120, 120, 120)
            self.cell(0, 8, "Lesson Plan Helper", align="L")
            self.ln(4)
            self.set_draw_color(226, 232, 240)
            self.line(10, self.get_y(), self.w - 10, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = LessonPDF(orientation="P", unit="mm", format="A4")
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    def _sanitize(text: str) -> str:
        """Replace common smart quotes/dashes and strip unencodable unicode."""
        replacements = {
            "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
            "\u2013": "-", "\u2014": "--", "\u2026": "...",
            "\u200b": "",  # zero-width space
            "\u2192": "->", "\u2190": "<-", "\u2022": "*",
            "\u00a0": " ",  # non-breaking space
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text.encode("latin-1", "ignore").decode("latin-1")

    # Sanitize the entire input first to be bulletproof
    md_text = _sanitize(md_text)

    # Process line by line
    for line in md_text.split("\n"):
        stripped = line.strip()

        # Heading level 1: # Title
        if stripped.startswith("# ") and not stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(30, 41, 59)
            pdf.ln(2)
            pdf.multi_cell(0, 8, _sanitize(stripped[2:].strip()))
            pdf.ln(3)

        # Heading level 2: ## Section
        elif stripped.startswith("## ") or (len(stripped) < 40 and any(stripped.lower().startswith(h) for h in ["objective", "essential question", "instructional materials", "teaching strategies", "hook", "direct instruction", "guided practice", "independent practice", "assessment", "assignments", "homework notes", "differentiation", "teacher reflection", "lesson plan:"]) and not stripped.startswith("#")):
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(37, 99, 235)
            pdf.ln(4)
            heading_text = stripped[3:].strip() if stripped.startswith("## ") else stripped
            pdf.multi_cell(0, 7, _sanitize(heading_text))
            pdf.ln(1)
            # Thin underline
            pdf.set_draw_color(226, 232, 240)
            pdf.line(10, pdf.get_y(), pdf.w - 10, pdf.get_y())
            pdf.ln(2)

        # Heading level 3: ### Subsection
        elif stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(51, 65, 85)
            pdf.ln(2)
            pdf.multi_cell(0, 6, _sanitize(stripped[4:].strip()))
            pdf.ln(1)

        # Bold metadata lines: **Key:** Value
        elif stripped.startswith("**"):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(51, 65, 85)
            # Parse bold parts
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            pdf.multi_cell(0, 5.5, _sanitize(clean))
            pdf.ln(0.5)

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(71, 85, 105)
            bullet_text = stripped[2:].strip()
            # Remove inline markdown bold/italic
            bullet_text = re.sub(r"\*\*(.+?)\*\*", r"\1", bullet_text)
            bullet_text = re.sub(r"\*(.+?)\*", r"\1", bullet_text)
            pdf.cell(6, 5.5, "")  # indent
            pdf.multi_cell(0, 5.5, _sanitize(f"  *  {bullet_text}"))
            pdf.ln(0.5)

        # Numbered items
        elif re.match(r"^\d+\.\s", stripped):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(71, 85, 105)
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            pdf.cell(4, 5.5, "")  # indent
            pdf.multi_cell(0, 5.5, _sanitize(f"  {clean}"))
            pdf.ln(0.5)

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            pdf.ln(2)
            pdf.set_draw_color(226, 232, 240)
            pdf.line(10, pdf.get_y(), pdf.w - 10, pdf.get_y())
            pdf.ln(3)

        # Empty line
        elif not stripped:
            pdf.ln(2)

        # Normal text
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(51, 65, 85)
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            pdf.multi_cell(0, 5.5, _sanitize(clean))
            pdf.ln(0.5)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ==============================================================================
# Markdown → Word DOCX  (via python-docx)
# ==============================================================================

def markdown_to_docx(md_text: str, meta: dict[str, Any] | None = None) -> bytes:
    """Convert a Markdown lesson plan to a Word .docx document.

    Parameters
    ----------
    md_text : str
        The lesson plan in Markdown format.
    meta : dict, optional
        Form metadata for document properties.

    Returns
    -------
    bytes
        The .docx file content.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 65, 85)

    # Configure heading styles
    for level, (size, color) in {
        "Heading 1": (Pt(18), RGBColor(30, 41, 59)),
        "Heading 2": (Pt(14), RGBColor(37, 99, 235)),
        "Heading 3": (Pt(12), RGBColor(51, 65, 85)),
    }.items():
        h_style = doc.styles[level]
        h_style.font.name = "Calibri"
        h_style.font.size = size
        h_style.font.color.rgb = color
        h_style.font.bold = True

    for line in md_text.split("\n"):
        stripped = line.strip()

        # Heading 1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:].strip(), level=1)

        # Heading 2
        elif stripped.startswith("## ") or (len(stripped) < 40 and any(stripped.lower().startswith(h) for h in ["objective", "essential question", "instructional materials", "teaching strategies", "hook", "direct instruction", "guided practice", "independent practice", "assessment", "assignments", "homework notes", "differentiation", "teacher reflection", "lesson plan:"]) and not stripped.startswith("#")):
            heading_text = stripped[3:].strip() if stripped.startswith("## ") else stripped
            doc.add_heading(heading_text, level=2)

        # Heading 3
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")

        # Numbered items
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            doc.add_paragraph(text, style="List Number")

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin border line via a run of underscores
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(226, 232, 240)
            run.font.size = Pt(8)

        # Bold metadata lines
        elif stripped.startswith("**"):
            p = doc.add_paragraph()
            # Parse **bold** segments
            parts = re.split(r"(\*\*.+?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Empty line → skip
        elif not stripped:
            continue

        # Normal paragraph
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==============================================================================
# Copy-to-Clipboard (JS-based for Streamlit)
# ==============================================================================

COPY_BUTTON_CSS = """
<style>
.copy-container {
    position: relative;
    display: inline-block;
    width: 100%;
}
.copy-btn {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    padding: 0.45rem 1rem;
    font-size: 0.85rem;
    font-weight: 500;
    font-family: 'Inter', system-ui, sans-serif;
    color: #334155;
    background: #FFFFFF;
    border: 1px solid #E2E8F0;
    border-radius: 6px;
    cursor: pointer;
    transition: all 0.15s ease;
    width: 100%;
    justify-content: center;
    box-sizing: border-box;
}
.copy-btn:hover {
    background: #F8FAFC;
    border-color: #CBD5E1;
}
.copy-btn.copied {
    color: #16A34A;
    border-color: #16A34A;
    background: #F0FDF4;
}
</style>
"""


def render_copy_button_html(text_to_copy: str, button_id: str = "copyBtn") -> str:
    """Generate an HTML/JS copy-to-clipboard button.

    Parameters
    ----------
    text_to_copy : str
        The plain text to copy.
    button_id : str
        Unique DOM id for the button.

    Returns
    -------
    str
        HTML string to render via st.markdown(unsafe_allow_html=True).
    """
    # Escape for JS string literal
    escaped = (
        text_to_copy
        .replace("\\", "\\\\")
        .replace("`", "\\`")
        .replace("$", "\\$")
    )

    return f"""{COPY_BUTTON_CSS}
<div class="copy-container">
<button class="copy-btn" id="{button_id}" onclick="
    navigator.clipboard.writeText(`{escaped}`).then(() => {{
        const btn = document.getElementById('{button_id}');
        btn.classList.add('copied');
        btn.innerHTML = '✅&nbsp; Copied to Clipboard';
        setTimeout(() => {{
            btn.classList.remove('copied');
            btn.innerHTML = '📋&nbsp; Copy to Clipboard';
        }}, 2000);
    }});
">
📋&nbsp; Copy to Clipboard
</button>
</div>"""
